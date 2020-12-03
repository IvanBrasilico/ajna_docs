from datetime import datetime

from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph

from ajna_commons.flask.log import logger


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def edit_text_tag(text: str, paragraph: Paragraph, conteudo: dict):
    inicio = text.find('{')
    fim = text.find('}')
    while inicio != -1:
        tag = text[inicio + 1:fim].strip()
        # print('*' + tag + '*')
        valor = conteudo.get(tag)
        if valor is not None:
            text = text[:inicio] + str(valor) + text[fim + 1:]
        else:
            text = text[:inicio] + f'** {tag} - vazio **' + text[fim + 1:]
        inicio = text.find('{')
        fim = text.find('}')
    paragraph.text = text


def edit_table_tag(text: str, paragraph: Paragraph, conteudo: dict, document: Document):
    tags = text[1:-1].split(':')
    # print(tags)
    valor = conteudo.get(tags[0])
    if valor is not None:
        paragraph.text = ' '
        table = document.add_table(rows=1, cols=len(tags) - 1)
        try:
            table.style = document.styles['Tabela']
        except Exception as err:
            logger.error(list(document.styles))
            logger.error(err, exc_info=True)
        move_table_after(table, paragraph)
        hdr_cells = table.rows[0].cells
        for ind_col, key in enumerate(tags[1:]):
            hdr_cells[ind_col].text = key.capitalize()
        for row in valor:
            row_cells = table.add_row().cells
            # print(row)
            for ind_col, key in enumerate(tags[1:]):
                content = row.get(key)
                row_cells[ind_col].text = str(content)


def edit_paragraph_tag(text: str, paragraph: Paragraph, conteudo: dict, document: Document):
    # print(f'*{text}*')
    tags = text[2:-2].split(':')
    # print(tags)
    valores = conteudo.get(tags[0])
    if valores is not None:
        paragraph.text = ' '
        for row in valores:
            # print(row)
            p = document.add_paragraph()
            paragraph._p.addnext(p._p)
            for key in tags[1:]:
                # print(key)
                if key.find(';') != -1:
                    titulo, key = key.split(';')
                else:
                    titulo = key.capitalize()
                valor = row[key]
                # print(key, valor)
                run = p.add_run('{}: {}'.format(titulo, valor))
                run.add_break()


def edit_image_tag(text: str, paragraph: Paragraph, conteudo: dict, document: Document):
    # print(f'*{text}*')
    tags = text[2:-2].split(':')
    # print(tags)
    valores = conteudo.get(tags[0])
    if valores is not None:
        paragraph.text = ' '
        for row in valores:
            # print(row)
            document.add_page_break()
            p = document.add_paragraph()
            for key in tags[1:]:
                # print(key)
                if key.find(';') != -1:
                    titulo, key = key.split(';')
                else:
                    titulo = key.capitalize()
                valor = row[key]
                # print(key, valor)
                run = p.add_run('{}: {}'.format(titulo, valor))
                run.add_break()
            document.add_picture(row['content'], width=Inches(5.5))


def paragraph_text_replace(paragraph: Paragraph, conteudo: dict, document: Document):
    text = paragraph.text.strip()
    if text and text.find('{{') != -1:
        edit_paragraph_tag(text, paragraph, conteudo, document)
    elif text and text.find('{') != -1:
        edit_text_tag(text, paragraph, conteudo)
    elif text and text.find('<<') != -1:
        edit_image_tag(text, paragraph, conteudo, document)
    elif text and text.find('<') != -1:
        edit_table_tag(text, paragraph, conteudo, document)



def docx_replacein(document: Document, conteudo: dict, user_name: str):
    agora = datetime.strftime(datetime.now(), '%d/%m/%Y %H:%M')
    footer = f'Emitido por {user_name} em {agora} pelo gerador de docx do sistema Ajna'
    section = document.sections[0]
    section.footer.paragraphs[0].text = footer
    for paragraph in section.header.paragraphs:
        try:
            paragraph_text_replace(paragraph, conteudo, document)
        except Exception as err:
            logger.error(err, exc_info=True)
    for table in section.header.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_text_replace(paragraph, conteudo, document)
    for paragraph in document.paragraphs:
        try:
            paragraph_text_replace(paragraph, conteudo, document)
        except Exception as err:
            logger.error(err, exc_info=True)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_text_replace(paragraph, conteudo, document)


def get_doc_generico_ovr(ovr: dict, documento: str, user_name: str):
    conteudo = {'unidade': 'ALFSTS', **ovr}
    document = Document(documento)
    docx_replacein(document, conteudo, user_name)
    return document
